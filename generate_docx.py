import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_blueprint_docx():
    doc = Document()

    # --- STYLE SETTINGS ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- TITLE PAGE ---
    for _ in range(5): doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("WHITEFLOWS MASTER BLUEPRINT")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(212, 168, 83) # Gold

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("The Definitive Technical & Operational Manual")
    run.font.size = Pt(16)
    run.italic = True

    for _ in range(3): doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version: 4.0 'Elite'\n")
    meta.add_run(f"Engineered by: Amburax\n")
    meta.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")

    doc.add_page_break()

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary & DNA', level=1)
    p = doc.add_paragraph()
    p.add_run("WhiteFlows represents a paradigm shift in financial technology specifically optimized for the Gujarat and international investment landscape. Traditionally, advisory platforms suffer from high maintenance and fragile database structures. WhiteFlows solves this with a ")
    p.add_run("Stateless / Resilient / High-Performance").bold = True
    p.add_run(" architecture.")

    p2 = doc.add_paragraph()
    p2.add_run("Built on the FastAPI engine, the platform is designed to handle high-frequency institutional enquiries with sub-500ms response times. The primary objective is absolute data integrity: ensuring that no investor lead is lost, even in the event of partial network failure.")

    # --- SECTION 2: FRONTEND ARCHITECTURE ---
    doc.add_heading('2. Client Experience & Security Envelope', level=1)
    
    doc.add_heading('2.1. Institutional Document Vault', level=2)
    p3 = doc.add_paragraph()
    p3.add_run("The entry point for HNI and Institutional clients is the Document Vault. This module uses advanced browser-side validation to ensure that all Know Your Customer (KYC) documents are processed before they even touch the server.")
    p3.add_run("\n\nKey features include:")
    
    bullets = [
        "Hardware-Accelerated File Scanning: Instant processing of PDFs and high-res images.",
        "3MB Security Envelope: Strict limit ensures that attachments are accepted by global SMTP relays (e.g., Gmail, Outlook).",
        "Base64 Transmission: Documents are encrypted into strings, avoiding common 'Multipart/Form-Data' vulnerabilities."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('2.2. Client-Side PDF Engineering', level=2)
    p4 = doc.add_paragraph()
    p4.add_run("To reduce server load and increase client satisfaction, WhiteFlows generates 'Wealth Receipts' instantly in the user's browser using the jsPDF engine. This provides the user with an immediate sense of authority and confirmation while the background tasks handle the mailing.")

    # --- SECTION 3: THE RESILIENCE CASCADE ---
    doc.add_page_break()
    doc.add_heading('3. The Backend Engine (The Fail-Safe)', level=1)
    
    doc.add_heading('3.1. 3-Tier Notification Cascade', level=2)
    p5 = doc.add_paragraph()
    p5.add_run("The core of WhiteFlows' reliability is the ")
    p5.add_run("Resilience Cascade").bold = True
    p5.add_run(". Communication is the lifeblood of a wealth management firm. If a lead notification fails, the firm loses money. Our engine prevents this with internal 'cascading' logic:")

    stats = doc.add_table(rows=4, cols=3)
    stats.style = 'Light Grid Accent 1'
    hdr = stats.rows[0].cells
    hdr[0].text = 'Priority'
    hdr[1].text = 'Service'
    hdr[2].text = 'Operational Role'
    
    row1 = stats.rows[1].cells
    row1[0].text = 'Primary'
    row1[1].text = 'Brevo SMTP'
    row1[2].text = 'High-volume professional relay for bulk enquiries.'
    
    row2 = stats.rows[2].cells
    row2[0].text = 'Secondary'
    row2[1].text = 'Brevo REST API'
    row2[2].text = 'HTTPS-based delivery if standard SMTP ports are blocked.'
    
    row3 = stats.rows[3].cells
    row3[0].text = 'Tertiary'
    row3[1].text = 'Gmail SMTP'
    row3[2].text = 'Bank-grade fallback for 100% notification certainty.'

    # --- SECTION 4: ADMIN COMMAND CENTER ---
    doc.add_heading('4. Business Intelligence & Control', level=1)
    p6 = doc.add_paragraph()
    p6.add_run("The Admin Command Center is a secure portal protected by JSON Web Tokens (JWT). It provides the leadership with a real-time view of firm growth.")
    
    bullets_admin = [
        "Regional Hotspot Mapping: Tagging leads by geography (Gujarat, UAE, Kenya, London) to identify trends.",
        "Momentum Tracking: Monitoring lead volume over a 24-hour window.",
        "Secure Management: One-click CSV exports and permanent record deletion protocol."
    ]
    for b in bullets_admin:
        doc.add_paragraph(b, style='List Bullet')

    # --- SECTION 5: UI/UX & MOBILE MASTERY (NEW) ---
    doc.add_heading('5. UI/UX Engineering & Mobile Mastery', level=1)
    p_ui = doc.add_paragraph()
    p_ui.add_run("A 'Wealth Management' platform must feel premium on every device. WhiteFlows is engineered with a high-contrast, gold-on-black aesthetic that remains responsive across all institutional hardware.")
    
    ui_features = [
        "Dark Mode Engineering: Reduced eye-strain for long-term data analysis.",
        "AOS Animation Engine: Subtle, professional transitions as the user scrolls.",
        "Liquid Grid Layout: The site adapts instantly from 4K workstations to ultra-high-resolution mobile devices."
    ]
    for feat in ui_features:
        doc.add_paragraph(feat, style='List Bullet')

    # --- SECTION 6: PERSISTENCE & BUSINESS CONTINUITY ---
    doc.add_page_break()
    doc.add_heading('6. Persistence & Data Sovereignty', level=1)
    p7 = doc.add_paragraph()
    p7.add_run("Your business data is the firm's most valuable asset. WhiteFlows protects it using professional infrastructure on ")
    p7.add_run("Render-managed Bare Metal").bold = True
    p7.add_run(".")

    doc.add_heading('6.1. The Persistent Disk', level=2)
    p8 = doc.add_paragraph()
    p8.add_run("We have engineered a 1GB Persistent Disk that follows the server. Even if the platform is updated or moved, the 'whiteflows.db' file remains untouched and secure on its own dedicated hardware.")

    doc.add_heading('6.2. Automated Backups & Schedulers', level=2)
    p9 = doc.add_paragraph()
    p9.add_run("Stability is monitored through three background processes:")
    
    scheds = [
        "3-Day Multi-Layer Backup: Full snapshots of the database are emailed to off-site vaults automatically.",
        "8:00 AM Intelligence Digest: Leadership receive a morning briefing of all client activity from yesterday.",
        "Rate-Limit Security: Automated protection against malicious bots or repeat spam attempts."
    ]
    for s in scheds:
        doc.add_paragraph(s, style='List Bullet')

    # --- SECTION 7: SYSTEM MAINTENANCE & READINESS (NEW) ---
    doc.add_heading('7. System Maintenance & Global Readiness', level=1)
    p_maint = doc.add_paragraph()
    p_maint.add_run("To ensure the platform remains 'Elite' for years, the following maintenance protocol and global features are integrated:")
    
    maint_bullets = [
        "Maintenance Checklist: Standardized weekly protocol for logs and disk audit.",
        "Multi-Region Readiness: Support for global phone codes (90+ countries) and multi-currency strategy displays.",
        "Scalability: Ready to handle thousands of concurrent enquiries with no hardware upgrades required."
    ]
    for m in maint_bullets:
        doc.add_paragraph(m, style='List Bullet')

    # --- FOOTER ---
    section = doc.sections[0]
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_p.add_run("© 2026 WhiteFlows International · Engineered by Amburax · Strictly Confidential")

    # --- SAVE ---
    save_path = "WHITEFLOWS_BLUEPRINT_FINAL.docx"
    doc.save(save_path)
    print(f"SUCCESS: {save_path} generated.")

if __name__ == "__main__":
    create_blueprint_docx()
